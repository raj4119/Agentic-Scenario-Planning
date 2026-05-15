"""Generates the CASCI Baseline Queries — Problem Statement Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "CASCI_Baseline_Query_Problems.docx")

PROBLEMS = [
    {
        "number": 1,
        "query": "Customer + Event Identification",
        "problem_statement": "What is the full context of the promotional event — which product, at which location, for which customer — and what did the planner originally intend to achieve?",
        "why_it_matters": (
            "Before any analysis can begin, the agent needs to know the exact scope of the event. "
            "Promotions in Blue Ridge span multiple product-location-customer combinations (tracked via "
            "event_prod_loc_cust_xref). Without this anchor, downstream queries have no join key and "
            "the agent cannot distinguish one event from another. This query also surfaces the planner's "
            "original intent — planned uplift, expected demand, and order quantity — which is the baseline "
            "against which AI recommendations are compared."
        ),
        "tables": ["event", "event_prod_loc_cust_xref", "product_location_customer_xref",
                   "product", "location", "customer", "product_class", "product_group",
                   "product_family", "segment"],
        "key_outputs": [
            "event_id, event_code, event_name, start_date, end_date, event_duration_days",
            "product_id, product_class, product_group, product_family",
            "location_id, customer_id, customer_segment",
            "planned_uplift_total, planned_demand_total, quantity_ordered",
        ],
        "agent": "Orchestrator / Demand Decomposition Agent",
        "decision_enabled": "Establishes the event scope and the planner's baseline expectation. Every other agent's output is contextualized against this.",
    },
    {
        "number": 2,
        "query": "Planned vs Actual Uplift",
        "problem_statement": "What did the planner expect demand to look like at T-14 (14 days before the event), and what actually happened week-by-week during the event window?",
        "why_it_matters": (
            "Promotional planning suffers from a common failure mode: the plan is set weeks in advance, "
            "but inventory decisions are made at T-14. By that point, the plan may already be stale. "
            "This query captures two views simultaneously: (1) the planned uplift as it existed in the "
            "system exactly 14 days before the event — using history tables to reconstruct the T-14 "
            "snapshot — and (2) the actual demand that materialized period-by-period from demand_history "
            "with signal='event'. The gap between these two views is where planning errors live. "
            "The Demand Decomposition Agent uses the actual branch to compute the pre-event baseline "
            "weekly average, which anchors all three lift scenarios."
        ),
        "tables": ["event_prod_loc_cust_uplift_History", "event_prod_loc_cust_xref_history",
                   "product_location_customer_xref", "demand_history", "demand_signal", "period"],
        "key_outputs": [
            "planned_uplift_units per period (from history at T-14)",
            "actual_event_demand per period (P1–P156 unpivoted)",
            "baseline_weekly_avg derived from pre-event periods",
            "planned vs actual gap visible for post-event retrospective",
        ],
        "agent": "Demand Decomposition Agent",
        "decision_enabled": "Provides the pre-event baseline weekly demand that all three scenario lift multipliers are applied to. Without this, the agent cannot compute Conservative / Moderate / Aggressive unit totals.",
    },
    {
        "number": 3,
        "query": "OOS + Overstock + Service Levels",
        "problem_statement": "Did this product experience stockouts or overstock in the 8 weeks before and 6 weeks after the event, and what were the fill rate and service level targets at each point in time?",
        "why_it_matters": (
            "Inventory health history is a leading indicator of planning quality. A product that was "
            "frequently out-of-stock before the event is a different risk profile than one that was "
            "overstocked — the first suggests demand is consistently under-estimated; the second suggests "
            "the planner over-ordered in a prior cycle. This query scans product_location_xref_history "
            "across a full ±8 week window, flags OOS and overstock conditions, and tracks how safety "
            "stock, fill rate goals, and lead times changed over time. The Supply Constraint Agent uses "
            "this to contextualize the ATP figure — 'is this a product that regularly runs lean?' — "
            "and to flag elevated stockout risk to the Scenario Generator."
        ),
        "tables": ["product_location_xref_history", "event_prod_loc_cust_xref",
                   "product_location_customer_xref", "event"],
        "key_outputs": [
            "inventory_BOH, inventory_onorder, inventory_backorder per snapshot",
            "available_to_promise (BOH - reserved + on_order)",
            "is_oos_flag (BOH <= out_of_stock_point)",
            "is_overstock_flag (overstock_units > 0)",
            "fill_rate_goal, service_level_objective at each snapshot",
            "lead_time_days over the window",
            "promo_window label: pre / during / post",
        ],
        "agent": "Supply Constraint Agent",
        "decision_enabled": "Tells the agent whether this product has a history of inventory problems that should raise or lower confidence in the pre-stage plan. A product with 3 OOS flags in the pre-window needs a larger safety buffer than one with none.",
    },
    {
        "number": 4,
        "query": "Comparable Event Matching + Lift Ratios",
        "problem_statement": "What do similar past promotions tell us about realistic lift ranges — and how confident should we be in those estimates?",
        "why_it_matters": (
            "This is the most analytically critical query in the workbook. The agent cannot invent lift "
            "assumptions from thin air — they must be grounded in historical evidence. This query "
            "implements a three-CTE pipeline: (1) anchor extracts the product group and customer segment "
            "of the target event; (2) comparable_events finds all past completed events in the same "
            "product group and customer segment; (3) comp_lift computes actual lift ratios for each "
            "comparable by dividing actual sales during the event by the pre-event 8-week baseline. "
            "PERCENTILE_CONT then derives p25 (Conservative), p50 (Moderate), and p75 (Aggressive) "
            "lift multipliers. The confidence tier (HIGH / MEDIUM / LOW) is set by comparable count: "
            ">=3 is HIGH, >=1 is MEDIUM, 0 is LOW with default fallback values. "
            "The LLM never computes these numbers — they arrive as pre-calculated facts."
        ),
        "tables": ["event", "event_prod_loc_cust_xref", "product_location_customer_xref",
                   "product", "product_group", "customer", "segment",
                   "demand_history", "demand_signal", "period"],
        "key_outputs": [
            "comparable_count and confidence_tier (HIGH / MEDIUM / LOW)",
            "conservative_lift_p25 — anchor for Conservative scenario",
            "moderate_lift_p50 — anchor for Moderate scenario (recommended default)",
            "aggressive_lift_p75 — anchor for Aggressive scenario",
            "actual_lift_ratio per comparable event",
            "pre_event_baseline, actual_sales_during, actual_event_uplift, lost_sales_during",
        ],
        "agent": "Demand Decomposition Agent",
        "decision_enabled": "Directly sets the three lift multipliers the Scenario Generator uses. Also sets the confidence score that flows through to the financial confidence interval. LOW confidence triggers a planner caveat in the final output.",
    },
    {
        "number": 5,
        "query": "Supply Constraint Snapshot at T-14",
        "problem_statement": "Exactly 14 days before the event, what was the inventory position and sourcing capacity — and was there enough time and stock to execute a pre-stage order for each scenario?",
        "why_it_matters": (
            "The most actionable output of the entire pipeline is the pre-stage order recommendation. "
            "But that recommendation is meaningless if the supply chain cannot physically execute it. "
            "This query reconstructs the state of the sourcing network as it existed at T-14 using "
            "three history tables: product_location_xref_history (inventory), "
            "product_sourcing_network_xref_History (order constraints), and sourcing_network_History "
            "(lead times). The feasibility flag is computed directly in SQL: if the remaining days "
            "until the event start are >= the supplier quoted lead time, pre-staging is feasible. "
            "If not, the agent must recommend ATP-only fulfillment. MOQ and max order quantity "
            "bound the orderable range for each scenario."
        ),
        "tables": ["event", "event_prod_loc_cust_xref", "product_location_customer_xref",
                   "product_location_xref_history", "product_sourcing_network_xref_History",
                   "sourcing_network", "sourcing_network_History"],
        "key_outputs": [
            "atp_at_t14 (BOH - reserved + on_order)",
            "pl_lead_time_days, sn_quoted_lead_time, total_lead_time",
            "moq (minimum order quantity), maximum_order_units",
            "days_until_event from T-14",
            "is_prestage_feasible (1 / 0)",
            "fulfillability_note — plain language label",
            "order_policy from sourcing network",
        ],
        "agent": "Supply Constraint Agent",
        "decision_enabled": "Determines which scenarios are physically executable. If Aggressive requires 3,000 units but max_order_qty is 2,000, the agent must mark it infeasible. This prevents the planner from being handed an impossible recommendation.",
    },
    {
        "number": 6,
        "query": "Financial Inputs at T-14",
        "problem_statement": "What are the unit economics — cost, price, carrying rate, and margin guardrails — needed to calculate the net financial impact of each scenario?",
        "why_it_matters": (
            "The Financial Impact Agent computes all monetary figures deterministically in Python — "
            "the LLM never generates a dollar amount. But the formulas need accurate inputs as of T-14, "
            "not today's prices. Prices and costs change over time; using the current value would "
            "misrepresent the trade-off the planner faced at decision time. This query joins "
            "product_location_xref_history at the T-14 snapshot to retrieve unit_cost, unit_price, "
            "and carrying_cost_rate. It also pulls segment-level margin guardrails (min_margin, "
            "max_margin from customer_segment) so the agent can flag if the recommended scenario "
            "would breach margin policy. Event duration (in weeks) is computed from the event dates "
            "to drive the carrying cost formula: pre_stage_qty × unit_cost × (rate/52) × rundown_weeks."
        ),
        "tables": ["event", "event_prod_loc_cust_xref", "product_location_customer_xref",
                   "product", "location", "customer", "customer_segment",
                   "product_location_xref_history"],
        "key_outputs": [
            "unit_cost, unit_price at T-14",
            "carrying_cost_rate (annual, e.g. 0.25 = 25%)",
            "fill_rate_goal, service_level_objective",
            "event_weeks (duration)",
            "min_margin, max_margin (segment guardrails)",
        ],
        "agent": "Financial Impact Agent",
        "decision_enabled": "Enables deterministic calculation of projected_revenue, carrying_cost, expected_stockout_cost, and net_financial_impact for each scenario. Margin guardrails let the agent flag a margin breach before the planner commits.",
    },
    {
        "number": 7,
        "query": "Cannibalization Candidates",
        "problem_statement": "Which substitute SKUs in the same product class, at the same location, are likely to lose sales when this product is promoted — and by how much?",
        "why_it_matters": (
            "Promotional lift is often partially an illusion: some of the 'extra' demand is customers "
            "switching from a similar product they would have bought anyway, not genuinely incremental "
            "volume. If the agent ignores cannibalization, it overstates net incremental demand and "
            "the planner over-orders. This query identifies substitute SKUs by matching on "
            "product_class_id at the same location, then computes each substitute's pre-event "
            "8-week baseline and their actual sales during the event window. The cannibalization_factor "
            "is expressed as (actual_during / pre_baseline) - 1.0 — a negative number. "
            "If no historical data exists for a substitute, a conservative default of -20% is applied. "
            "The Demand Decomposition Agent surfaces these factors; the Scenario Generator incorporates "
            "them into the net incremental demand calculation and the final narrative."
        ),
        "tables": ["event", "event_prod_loc_cust_xref", "product_location_customer_xref",
                   "product", "product_class", "demand_history", "demand_signal", "period"],
        "key_outputs": [
            "sub_product_id, sub_product_code, sub_product_name",
            "shared_product_class",
            "sub_pre_event_baseline (8-week avg before event)",
            "sub_actual_sales_during (sales during promotion window)",
            "cannibalization_factor: (actual/baseline) - 1.0, default -0.20",
        ],
        "agent": "Demand Decomposition Agent",
        "decision_enabled": "Adjusts net incremental demand downward to account for intra-category switching. Prevents over-ordering by ensuring the planner sees true incremental volume, not gross lift.",
    },
]


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1, color="1F3864"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = RGBColor.from_string(color)
    return h


def add_label_value(doc, label, value, bold_label=True):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = bold_label
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_bullet_list(doc, items, label=None):
    if label:
        p = doc.add_paragraph()
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(10.5)


def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("CASCI — Baseline Query Problem Statements", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor.from_string("1F3864")
        run.font.size = Pt(20)

    subtitle = doc.add_paragraph(
        "Composable Agentic Supply Chain Intelligence — Promotional Event Scenario Planning\n"
        "7 Foundational Data Problems & the Queries That Solve Them"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string("404040")
    doc.add_paragraph()

    # Overview table
    add_heading(doc, "Overview", level=1)
    overview_p = doc.add_paragraph(
        "Each of the 7 queries in the Baseline Events workbook solves a distinct analytical problem "
        "that a human planner would have to answer manually before making a pre-stage inventory decision. "
        "Taken together, they give the CASCI agent the complete picture it needs to generate "
        "Conservative, Moderate, and Aggressive scenarios with grounded demand estimates, "
        "supply feasibility checks, and deterministic financial impacts."
    )
    overview_p.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    # Summary table
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(["#", "Query Name", "Agent", "Core Question"]):
        hdr[i].text = h
        set_cell_bg(hdr[i], "1F3864")
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True
                run.font.size = Pt(10)

    alt_colors = ["FFFFFF", "EBF3FB"]
    for i, p in enumerate(PROBLEMS):
        row = table.add_row().cells
        row[0].text = str(p["number"])
        row[1].text = p["query"]
        row[2].text = p["agent"].split("/")[0].strip()
        row[3].text = p["problem_statement"][:80] + "..."
        bg = alt_colors[i % 2]
        for cell in row:
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)

    doc.add_page_break()

    # Detail sections
    for prob in PROBLEMS:
        # Problem header bar
        add_heading(doc, f"Query {prob['number']} — {prob['query']}", level=1, color="1F3864")

        # Problem statement box
        ps_para = doc.add_paragraph()
        ps_para.paragraph_format.left_indent = Inches(0.2)
        ps_para.paragraph_format.space_before = Pt(4)
        ps_para.paragraph_format.space_after = Pt(8)
        r = ps_para.add_run("Problem Statement:  ")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor.from_string("C00000")
        r2 = ps_para.add_run(prob["problem_statement"])
        r2.font.size = Pt(11)
        r2.italic = True

        # Why it matters
        add_heading(doc, "Why This Problem Matters", level=2, color="2E75B6")
        wim = doc.add_paragraph(prob["why_it_matters"])
        wim.runs[0].font.size = Pt(11)
        wim.paragraph_format.space_after = Pt(8)

        # Source tables
        add_bullet_list(doc, prob["tables"], label="Source Tables:")
        doc.add_paragraph()

        # Key outputs
        add_bullet_list(doc, prob["key_outputs"], label="Key Output Columns:")
        doc.add_paragraph()

        # Agent + Decision
        add_label_value(doc, "Consumed by Agent", prob["agent"])
        add_label_value(doc, "Decision Enabled", prob["decision_enabled"])

        if prob["number"] < 7:
            doc.add_page_break()
        else:
            doc.add_paragraph()

    # Footer note
    doc.add_paragraph()
    footer_p = doc.add_paragraph(
        "All 7 queries are parameterised on event_id. The Baseline Events workbook contains "
        "versions filtered to event 41 for validation. In the CASCI pipeline, each query is "
        "wrapped as a Python tool function and called at runtime with the target event_id."
    )
    footer_p.runs[0].font.size = Pt(10)
    footer_p.runs[0].font.italic = True
    footer_p.runs[0].font.color.rgb = RGBColor.from_string("595959")

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_doc()
